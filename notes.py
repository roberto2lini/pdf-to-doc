from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt

# Load the PDF file
pdf_path = "/home/ladmin/Class10.w2v.pdf"
pdf = PdfReader(open(pdf_path, "rb"))

# Create a new Document
doc = Document()

# Process each page of the PDF
for page_num in range(len(pdf.pages)):
    # Extract the text from the current page
    page = pdf.pages[page_num]
    text = page.extract_text()

    # Split the text into lines
    lines = text.split('\n')

    # Add the text to the Word document excluding the line with the specified content
    for line in lines:
        if "G. Cassani & A. AlishahiCL, Class # 10: Word2Vec" in line:
            continue  # Skip this line if it contains the excluded content
        # Add the line to the Word document
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(12)
    
    # Add a paragraph for notes
    doc.add_paragraph("Notes:", style='Normal')
    doc.add_paragraph()

    # Add a page break after each slide
    if page_num < len(pdf.pages) - 1:
        doc.add_page_break()

# Save the DOCX file
docx_path = "/home/ladmin/Class10_Notes.docx"
doc.save(docx_path)

docx_path